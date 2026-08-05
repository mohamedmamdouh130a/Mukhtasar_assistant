import os
import streamlit as st
import requests
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document as DocxDocument
from fpdf import FPDF
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from groq import Groq
import arabic_reshaper
from bidi.algorithm import get_display

st.set_page_config(page_title="Mukhtasar", page_icon="🎯", layout="centered")
st.title("🎯 Mukhtasar")
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@300;400;500;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Cairo', sans-serif !important;
        font-size: 19px !important;
    }

    h2 { font-size: 2.0rem !important; }
    h3 { font-size: 1.5rem !important; }
    .subtitle {
        color: #94A3B8;
        font-size: 15px !important;
        margin-top: -10px;
        margin-bottom: 20px;
    }
    </style>
""", unsafe_allow_html=True)

api_key = st.secrets.get("GROQ_API_KEY")
client = Groq(api_key=api_key)

def ask_groq(prompt, lang):
    try:
        res = client.chat.completions.create(
            messages=[
                {"role": "system", "content": f"You are a strict summarization and QA assistant. Answer entirely in {lang}. Rely ONLY and STRICTLY on the provided Context. Do NOT hallucinate, assume, or invent any external facts, names, or events not present in the text."},
                {"role": "user", "content": prompt}
            ],
            model="llama-3.3-70b-versatile",
            temperature=0.1
        )
        return res.choices[0].message.content
    except Exception as e:
        return f"Error communicating with Groq API: {e}"

@st.cache_resource
def get_embeddings():
    return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

def process_text(text):
    chunks = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100).split_documents([Document(page_content=text)])
    return FAISS.from_documents(chunks, get_embeddings())

def fix_arabic(text):
    if not isinstance(text, str):
        return str(text)
    try:
        if any("\u0600" <= c <= "\u06ff" for c in text):
            reshaped_text = arabic_reshaper.reshape(text)
            return get_display(reshaped_text)
        return text
    except Exception:
        return text

class CairoPDF(FPDF):
    def __init__(self):
        super().__init__()
        font_path = "Cairo-Regular.ttf"
        if os.path.exists(font_path):
            self.add_font("Cairo", "", font_path, uni=True)
        else:
            self.add_font("Cairo", "", "Cairo-Regular.ttf", uni=True)

def export_docx(summary, history):
    doc = DocxDocument()
    doc.add_heading('Mukhtasar Report', 0)
    doc.add_heading('Summary:', level=1)
    doc.add_paragraph(summary)
    doc.add_heading('Chat History:', level=1)
    for m in history:
        doc.add_paragraph(f"{'User' if m['role']=='user' else 'Assistant'}: {m['content']}")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_pdf(summary, history):
    pdf = CairoPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)    
    font_name = "Cairo"
    try:
        pdf.set_font(font_name, size=12)
    except Exception:
        font_name = "Helvetica"
        pdf.set_font(font_name, size=12)

    pdf.cell(200, 10, txt=fix_arabic("Mukhtasar Report"), ln=True, align='C')
    pdf.ln(5)
    
    pdf.set_font(font_name, size=11)
    pdf.cell(200, 8, txt=fix_arabic("Summary:"), ln=True)
    
    pdf.set_font(font_name, size=10)
    pdf.multi_cell(190, 6, txt=fix_arabic(summary))
    pdf.ln(5)
    
    pdf.set_font(font_name, size=11)
    pdf.cell(200, 8, txt=fix_arabic("Chat History:"), ln=True)
    
    pdf.set_font(font_name, size=10)
    for m in history:
        role = "User: " if m['role']=='user' else "Assistant: "
        content_text = f"{role}{m['content']}"
        pdf.multi_cell(190, 6, txt=fix_arabic(content_text))
        pdf.ln(2)
        
    return bytes(pdf.output())

st.markdown('<p class="subtitle">AI assistant for Summarizing text, URLs and PDFs with interactive chat.</p>', unsafe_allow_html=True)
lang = st.selectbox("Language:", ["Arabic", "English", "French"])
source = st.selectbox("Choose Source Type", ["URL", "PDF", "Text"])

if "last_source" not in st.session_state:
    st.session_state.last_source = source

if st.session_state.last_source != source:
    st.session_state.last_source = source
    st.session_state.raw_text = ""
    st.session_state.pop("vectordb", None)
    st.session_state.pop("summary", None)
    st.session_state.messages = []

if "raw_text" not in st.session_state:
    st.session_state.raw_text = ""

if source == "URL":
    url = st.text_input("Paste URL:")
    if url and st.button("Process"):
        res = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'}, timeout=10)
        soup = BeautifulSoup(res.text, 'html.parser')
        for s in soup(["script", "style", "nav", "footer"]):
            s.decompose()
        st.session_state.raw_text = " ".join(soup.get_text().split())
        st.session_state.pop("vectordb", None)

elif source == "PDF":
    pdf_file = st.file_uploader("Upload PDF", type=["pdf"])
    if pdf_file and st.button("Process"):
        with open("tmp.pdf", "wb") as f: f.write(pdf_file.getbuffer())
        st.session_state.raw_text = "\n".join([p.page_content for p in PyPDFLoader("tmp.pdf").load()])
        os.remove("tmp.pdf")
        st.session_state.pop("vectordb", None)

elif source == "Text":
    text_input = st.text_area("Paste Text:")
    if text_input and st.button("Process"):
        st.session_state.raw_text = text_input
        st.session_state.pop("vectordb", None)

if st.session_state.raw_text and "vectordb" not in st.session_state:
    with st.spinner("Processing..."):
        st.session_state.vectordb = process_text(st.session_state.raw_text)
        st.session_state.summary = ask_groq(f"Provide a concise, strict summary based ONLY on this text:\n\n{st.session_state.raw_text[:3000]}", lang)
        st.session_state.messages = []

if "summary" in st.session_state and st.session_state.summary:
    st.markdown("-----")
    col_s1, col_s2 = st.columns([0.85, 0.15])
    with col_s1:
        st.subheader("Summary:")
    with col_s2:
        st.copy_button("Copy", st.session_state.summary)
    st.write(st.session_state.summary)

if st.session_state.get("vectordb"):
    st.markdown("-----")
    st.subheader("Chat")
    
    if "messages" not in st.session_state:
        st.session_state.messages = []

    for idx, m in enumerate(st.session_state.messages):
        with st.chat_message(m["role"]):
            st.markdown(m["content"])
            st.copy_button("Copy", m["content"], key=f"copy_msg_{idx}")

    if q := st.chat_input("Ask a question..."):
        st.session_state.messages.append({"role": "user", "content": q})
        with st.chat_message("user"): st.markdown(q)
        
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                docs = st.session_state.vectordb.similarity_search(q, k=2)
                ctx = "\n\n".join([d.page_content for d in docs])
                ans = ask_groq(f"Answer the question strictly and solely based on the provided Context. Do not invent any outside facts.\n\nContext:\n{ctx}\n\nQuestion: {q}\nAnswer:", lang)
                st.markdown(ans)
        st.session_state.messages.append({"role": "assistant", "content": ans})

    if st.session_state.messages:
        st.markdown("---")
        st.subheader("Export Report")
        c1, c2 = st.columns(2)
        with c1:
            st.download_button("Download Word", export_docx(st.session_state.summary, st.session_state.messages), "Mukhtasar.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with c2:
            st.download_button("Download PDF", export_pdf(st.session_state.summary, st.session_state.messages), "Mukhtasar.pdf", "application/pdf")
